from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_MD = DOCS / "AGLI_Manual_Book.md"
OUT_DOCX = DOCS / "AGLI_Manual_Book.docx"


manual = {
    "English": {
        "subtitle": "Complete user manual for Amsterdam Game Lab Intelligence",
        "intro": [
            "AGLI is a lightweight content-intelligence workspace for Amsterdam Game Lab. It helps a small team research content that already performs well, identify repeatable patterns, and turn those patterns into credible draft content for Pro Actief.",
            "The application is designed for HR, L&D, team leads, managers, and internal content staff. It does not auto-publish. Every metric, claim, quote, and client reference must be checked by a human before use.",
        ],
        "principles": [
            "Research first: import or review evidence before writing.",
            "No fabricated claims: use verified platform exports, real quotes, and approved client references only.",
            "Drafts are starting points: AGLI writes editable ideas, not final legal or brand-approved copy.",
            "Scenario-driven output: changing ICP, platform, objective, tone, or selected row changes the draft direction.",
            "Small-team workflow: keep research, patterns, calendar, library, reports, and team ownership in one place.",
        ],
        "quick_start": [
            "Open the landing page and press Login.",
            "Use the demo credentials or edit the profile details, then press Login. After login, AGLI opens Overview.",
            "Go to Research, select ICP, platform, objective, and tone, then review ranked content rows.",
            "Import CSV or use the demo scenario rows for a presentation flow.",
            "Click a strong row, open Draft Studio, select a tab, and press Generate draft.",
            "Review the draft, save it, add it to Calendar, or export Reports for the hackathon deck.",
        ],
        "global": [
            ["Sidebar / mobile hamburger", "Opens main navigation. On mobile, the sidebar becomes a hamburger drawer."],
            ["Top verified-metrics message", "Reminds users that imported metrics and claims need human verification."],
            ["Help icon", "Opens quick workflow guidance."],
            ["Notification icon", "Shows current workspace alerts and data status."],
            ["Profile menu", "Opens account actions: Update profile, Data sources, Shortcuts, and Logout."],
            ["Dropdowns", "Open with the arrow, close by clicking anywhere outside, and update the active scenario."],
            ["Language setting", "Settings can switch the interface between English, Indonesian, and Dutch."],
        ],
        "pages": [
            ["Overview", "Main landing page after login. Shows the content workflow, key metrics, current scenario status, and quick actions.", "Use Open Research, Import CSV, Open Pattern Engine, Review Calendar, and report shortcuts to start the demo."],
            ["Research", "Core workbench for ranking content and drafting. Filters reshape the table and Draft Studio output.", "Use ICP, Platform, Objective, Tone, Date Range, Filters, Save Search, Import CSV, table row selection, pagination, Pattern Insight links, Draft Studio tabs, Generate Draft, Save Draft, and Add to Calendar."],
            ["Pattern Engine", "Turns high-performing rows into repeatable patterns.", "Review hook types, topics, formats, ideal lengths, visual styles, and evidence rows. Use it after data import or scenario selection."],
            ["Calendar", "Plans reviewed draft ideas into a publishing schedule.", "Check status, platform, dates, and titles. Add items from Draft Studio for review instead of auto-publishing."],
            ["Content Library", "Stores saved drafts, outlines, and reusable assets.", "Review status, audience, content type, and next action. Saved drafts from Research appear here."],
            ["Reports", "Summarizes performance signals for presentation and submission.", "Compare platforms, median engagement, sample size, best signal, coverage, and export the report/evidence pack."],
            ["Saved Searches", "Stores reusable research scenarios.", "Use saved combinations of ICP, platform, objective, and tone to repeat a discovery flow."],
            ["Brand Voice", "Holds writing guardrails for Pro Actief and Amsterdam Game Lab.", "Review credibility rules, tone boundaries, proof requirements, and content principles before publishing."],
            ["Team", "Shows ownership and review responsibilities.", "Use it to explain who researches, writes, reviews, and approves content."],
            ["Settings", "Controls language, storage, publishing safety, and backup.", "Switch language, toggle verification safeguards, export backup, and clear local workspace data when needed."],
            ["Update Profile", "Edits the workspace profile used in the navbar and profile menu.", "Change name, role, email, and profile photo for demo or presentation personalization."],
            ["Data Sources", "Explains and manages imported evidence.", "Import CSV, review accepted source types, reset demo scenario, and export backup data."],
            ["Shortcuts", "Provides quick routes for common workflows.", "Jump to Research, Content Library, Reports, or repeat the demo flow."],
        ],
        "buttons": [
            ["Login", "Starts the app and opens Overview."],
            ["Back to landing", "Returns from the login form to the landing page."],
            ["Import CSV", "Adds research rows from platform exports or manual research."],
            ["Filters", "Opens the active scenario filter drawer."],
            ["Save search", "Stores the current filter combination."],
            ["Generate draft", "Creates a draft from active patterns, selected row, ICP, platform, objective, and tone."],
            ["Save draft", "Stores the current draft in Content Library."],
            ["Add to Calendar", "Creates a planning item for review."],
            ["Export analysis / Export report / Export backup", "Downloads JSON evidence for reproducibility and hackathon submission."],
            ["Reset demo scenario", "Restores the built-in demo dataset for presentation use."],
            ["Clear workspace", "Removes local browser data for a clean demo."],
            ["Logout", "Returns to the landing page."],
        ],
        "workflow": [
            "Content research: choose a scenario, import evidence, rank rows, and select a strong reference item.",
            "Pattern extraction: AGLI summarizes hook, topic, format, length, and visual-style patterns from the active evidence set.",
            "Draft creation: Draft Studio combines the selected pattern with Pro Actief guardrails and produces copy, hooks, video outlines, or hashtags.",
            "Human review: the team checks factual accuracy, brand fit, legal risk, and proof quality.",
            "Planning and reporting: save draft, add calendar item, export reports, and document what was verified.",
        ],
        "csv": [
            ["title", "Content title or hook."],
            ["platform", "LinkedIn, YouTube, Instagram, or TikTok."],
            ["account", "Source account name."],
            ["type", "Text post, carousel, video, article, or similar format."],
            ["views", "Verified or imported view count."],
            ["engagements", "Likes, comments, shares, or platform engagement count."],
            ["date", "Publish date."],
            ["hookType / topic / visualStyle", "Optional pattern fields that improve Draft Studio output."],
            ["verificationNote", "Optional human note about source quality."],
        ],
        "demo_script": [
            "Start on the landing page and explain the problem: serious games are category-creation, so content must earn attention.",
            "Login and show Overview as the main command center.",
            "Open Research, adjust the scenario dropdowns, and select a high-performing row.",
            "Show Pattern Insights and explain how the app identifies hooks, topics, formats, lengths, and visual style.",
            "Generate a Draft Studio output, then save it and add it to Calendar.",
            "Open Reports and export the evidence pack.",
            "Close with guardrails: no auto-publishing, no fabricated claims, and every draft requires human review.",
        ],
        "troubleshooting": [
            ["A page changes after refresh", "Refresh now keeps the current page. Login is the only flow that intentionally opens Overview."],
            ["Dropdown stays open", "Click anywhere outside the dropdown to close it."],
            ["Drawer stays open", "Click outside the right drawer or use the X icon."],
            ["Draft feels wrong", "Change ICP, platform, objective, tone, selected row, or content pillar, then generate again."],
            ["Metrics are not verified", "Import official platform exports or add verification notes before using the content externally."],
        ],
    },
    "Indonesia": {
        "subtitle": "Panduan lengkap pengguna Amsterdam Game Lab Intelligence",
        "intro": [
            "AGLI adalah workspace content-intelligence ringan untuk Amsterdam Game Lab. Alat ini membantu tim kecil meneliti konten yang performanya bagus, menemukan pola yang bisa diulang, lalu mengubah pola tersebut menjadi draft konten kredibel untuk Pro Actief.",
            "Aplikasi ini dibuat untuk HR, L&D, team lead, manager, dan staf konten internal. AGLI tidak melakukan auto-publish. Semua metrik, klaim, kutipan, dan referensi klien harus dicek manusia sebelum digunakan.",
        ],
        "principles": [
            "Riset dulu: impor atau tinjau bukti sebelum menulis.",
            "Tidak ada klaim palsu: gunakan export platform, kutipan asli, dan referensi klien yang disetujui.",
            "Draft adalah awal: AGLI membuat ide yang bisa diedit, bukan copy final yang sudah legal atau brand approved.",
            "Output berbasis skenario: perubahan ICP, platform, objective, tone, atau baris referensi akan mengubah arah draft.",
            "Alur tim kecil: riset, pola, kalender, library, laporan, dan ownership ada dalam satu tempat.",
        ],
        "quick_start": [
            "Buka landing page lalu tekan Login.",
            "Gunakan akun demo atau ubah detail profil, lalu tekan Login. Setelah login, AGLI masuk ke Overview.",
            "Buka Research, pilih ICP, platform, objective, dan tone, lalu tinjau konten yang sudah diranking.",
            "Import CSV atau gunakan data skenario demo untuk alur presentasi.",
            "Klik row yang kuat, buka Draft Studio, pilih tab, lalu tekan Generate draft.",
            "Review draft, simpan, masukkan ke Calendar, atau export Reports untuk deck hackathon.",
        ],
        "global": [
            ["Sidebar / hamburger mobile", "Membuka navigasi utama. Di mobile, sidebar berubah menjadi drawer hamburger."],
            ["Pesan verified metrics", "Mengingatkan bahwa metrik dan klaim impor harus diverifikasi manusia."],
            ["Ikon Help", "Membuka panduan workflow cepat."],
            ["Ikon Notification", "Menampilkan alert workspace dan status data."],
            ["Menu profil", "Membuka Update profile, Data sources, Shortcuts, dan Logout."],
            ["Dropdown", "Dibuka lewat anak panah, ditutup dengan klik area luar, dan mengubah skenario aktif."],
            ["Language setting", "Settings dapat mengganti bahasa interface ke Inggris, Indonesia, atau Belanda."],
        ],
        "pages": [
            ["Overview", "Halaman utama setelah login. Menampilkan workflow, metrik ringkas, status skenario, dan quick action.", "Gunakan Open Research, Import CSV, Open Pattern Engine, Review Calendar, dan shortcut report untuk memulai demo."],
            ["Research", "Workbench utama untuk ranking konten dan drafting. Filter mengubah tabel dan hasil Draft Studio.", "Gunakan ICP, Platform, Objective, Tone, Date Range, Filters, Save Search, Import CSV, pilihan row tabel, pagination, Pattern Insight links, tab Draft Studio, Generate Draft, Save Draft, dan Add to Calendar."],
            ["Pattern Engine", "Mengubah row performa tinggi menjadi pola yang bisa diulang.", "Tinjau hook, topik, format, panjang ideal, visual style, dan evidence row setelah import data atau memilih skenario."],
            ["Calendar", "Merencanakan ide draft yang sudah direview ke jadwal publikasi.", "Cek status, platform, tanggal, dan judul. Item dari Draft Studio masuk untuk review, bukan auto-publish."],
            ["Content Library", "Menyimpan draft, outline, dan aset reusable.", "Tinjau status, audience, tipe konten, dan next action. Draft yang disimpan dari Research muncul di sini."],
            ["Reports", "Merangkum sinyal performa untuk presentasi dan submission.", "Bandingkan platform, median engagement, sample size, best signal, coverage, dan export evidence pack."],
            ["Saved Searches", "Menyimpan skenario riset yang bisa dipakai ulang.", "Gunakan kombinasi ICP, platform, objective, dan tone untuk mengulang flow discovery."],
            ["Brand Voice", "Menyimpan guardrail penulisan untuk Pro Actief dan Amsterdam Game Lab.", "Review aturan kredibilitas, batas tone, kebutuhan bukti, dan prinsip konten sebelum publish."],
            ["Team", "Menampilkan ownership dan tanggung jawab review.", "Gunakan untuk menjelaskan siapa yang riset, menulis, review, dan approve konten."],
            ["Settings", "Mengatur bahasa, storage, keamanan publishing, dan backup.", "Ganti bahasa, nyalakan safeguard verifikasi, export backup, dan clear workspace lokal bila perlu."],
            ["Update Profile", "Mengedit profil workspace yang tampil di navbar dan menu profil.", "Ubah nama, role, email, dan foto profil untuk demo atau personalisasi presentasi."],
            ["Data Sources", "Menjelaskan dan mengelola bukti yang diimpor.", "Import CSV, cek sumber yang diterima, reset demo scenario, dan export backup data."],
            ["Shortcuts", "Memberi akses cepat ke workflow umum.", "Lompat ke Research, Content Library, Reports, atau ulangi demo flow."],
        ],
        "buttons": [
            ["Login", "Masuk ke aplikasi dan membuka Overview."],
            ["Back to landing", "Kembali dari form login ke landing page."],
            ["Import CSV", "Menambahkan row riset dari export platform atau riset manual."],
            ["Filters", "Membuka drawer filter skenario aktif."],
            ["Save search", "Menyimpan kombinasi filter saat ini."],
            ["Generate draft", "Membuat draft dari pola aktif, row terpilih, ICP, platform, objective, dan tone."],
            ["Save draft", "Menyimpan draft saat ini ke Content Library."],
            ["Add to Calendar", "Membuat item planning untuk review."],
            ["Export analysis / Export report / Export backup", "Mengunduh JSON evidence untuk reproducibility dan submission hackathon."],
            ["Reset demo scenario", "Mengembalikan dataset demo bawaan untuk presentasi."],
            ["Clear workspace", "Menghapus data browser lokal untuk demo bersih."],
            ["Logout", "Keluar dan kembali ke landing page."],
        ],
        "workflow": [
            "Riset konten: pilih skenario, import bukti, ranking row, dan pilih referensi yang kuat.",
            "Ekstraksi pola: AGLI merangkum pola hook, topik, format, durasi, dan visual style dari evidence aktif.",
            "Pembuatan draft: Draft Studio menggabungkan pola terpilih dengan guardrail Pro Actief untuk membuat copy, hooks, video outline, atau hashtags.",
            "Review manusia: tim mengecek akurasi fakta, brand fit, risiko legal, dan kualitas bukti.",
            "Planning dan reporting: save draft, add calendar item, export reports, dan dokumentasikan bukti yang diverifikasi.",
        ],
        "csv": [
            ["title", "Judul konten atau hook."],
            ["platform", "LinkedIn, YouTube, Instagram, atau TikTok."],
            ["account", "Nama akun sumber."],
            ["type", "Text post, carousel, video, article, atau format serupa."],
            ["views", "Jumlah view yang diverifikasi atau diimpor."],
            ["engagements", "Likes, comments, shares, atau total engagement platform."],
            ["date", "Tanggal publish."],
            ["hookType / topic / visualStyle", "Field pola opsional yang membuat output Draft Studio lebih akurat."],
            ["verificationNote", "Catatan manusia tentang kualitas sumber."],
        ],
        "demo_script": [
            "Mulai dari landing page dan jelaskan masalah: serious games adalah category-creation, jadi konten harus bisa menarik perhatian.",
            "Login dan tampilkan Overview sebagai command center.",
            "Buka Research, ubah dropdown skenario, dan pilih row performa tinggi.",
            "Tampilkan Pattern Insights dan jelaskan bagaimana app menemukan hook, topic, format, length, dan visual style.",
            "Generate output Draft Studio, lalu save dan add to Calendar.",
            "Buka Reports dan export evidence pack.",
            "Tutup dengan guardrail: tidak auto-publish, tidak ada klaim palsu, dan semua draft harus direview manusia.",
        ],
        "troubleshooting": [
            ["Halaman berubah setelah refresh", "Refresh sekarang tetap di halaman yang sedang dibuka. Hanya login yang sengaja membuka Overview."],
            ["Dropdown tidak tertutup", "Klik area luar dropdown untuk menutupnya."],
            ["Drawer tidak tertutup", "Klik area luar drawer kanan atau gunakan ikon X."],
            ["Draft kurang cocok", "Ubah ICP, platform, objective, tone, row terpilih, atau content pillar, lalu generate ulang."],
            ["Metrik belum verified", "Import export resmi platform atau tambahkan verification note sebelum dipakai eksternal."],
        ],
    },
    "Nederlands": {
        "subtitle": "Volledige gebruikershandleiding voor Amsterdam Game Lab Intelligence",
        "intro": [
            "AGLI is een lichte content-intelligence workspace voor Amsterdam Game Lab. De tool helpt een klein team content te onderzoeken die al goed presteert, herhaalbare patronen te vinden, en die patronen om te zetten naar geloofwaardige conceptcontent voor Pro Actief.",
            "De applicatie is bedoeld voor HR, L&D, teamleiders, managers en interne contentmedewerkers. AGLI publiceert niets automatisch. Elke metric, claim, quote en klantreferentie moet door een mens worden gecontroleerd voor gebruik.",
        ],
        "principles": [
            "Eerst onderzoek: importeer of controleer bewijs voordat je schrijft.",
            "Geen verzonnen claims: gebruik alleen geverifieerde platformexports, echte quotes en goedgekeurde klantreferenties.",
            "Concepten zijn startpunten: AGLI schrijft bewerkbare ideeen, geen definitieve juridische of brand-approved copy.",
            "Scenario-gestuurde output: ICP, platform, objective, tone en geselecteerde rij veranderen de richting van het concept.",
            "Workflow voor kleine teams: onderzoek, patronen, kalender, library, rapporten en ownership staan op een plek.",
        ],
        "quick_start": [
            "Open de landing page en druk op Login.",
            "Gebruik het demo-account of pas profielgegevens aan, en druk op Login. Na login opent AGLI Overview.",
            "Ga naar Research, kies ICP, platform, objective en tone, en bekijk de gerangschikte content.",
            "Importeer CSV of gebruik de demo-scenario's voor een presentatieflow.",
            "Klik op een sterke rij, open Draft Studio, kies een tab en druk op Generate draft.",
            "Controleer het concept, sla het op, voeg het toe aan Calendar, of exporteer Reports voor het hackathon deck.",
        ],
        "global": [
            ["Sidebar / mobiele hamburger", "Opent de hoofdnavigatie. Op mobiel wordt de sidebar een hamburger drawer."],
            ["Verified-metrics bericht", "Herinnert gebruikers eraan dat metrics en claims door mensen moeten worden geverifieerd."],
            ["Help-icoon", "Opent korte workflowhulp."],
            ["Notification-icoon", "Toont workspace-alerts en datastatus."],
            ["Profielmenu", "Opent Update profile, Data sources, Shortcuts en Logout."],
            ["Dropdowns", "Open met de pijl, sluit door buiten de dropdown te klikken, en wijzig het actieve scenario."],
            ["Taalinstelling", "Settings kan de interface wisselen tussen Engels, Indonesisch en Nederlands."],
        ],
        "pages": [
            ["Overview", "Hoofdpagina na login. Toont de contentworkflow, kernmetrics, scenariostatus en snelle acties.", "Gebruik Open Research, Import CSV, Open Pattern Engine, Review Calendar en rapportshortcuts om de demo te starten."],
            ["Research", "Belangrijkste werkbank voor contentranking en drafting. Filters veranderen de tabel en Draft Studio output.", "Gebruik ICP, Platform, Objective, Tone, Date Range, Filters, Save Search, Import CSV, tabelrijselectie, pagination, Pattern Insight links, Draft Studio tabs, Generate Draft, Save Draft en Add to Calendar."],
            ["Pattern Engine", "Zet goed presterende rijen om in herhaalbare patronen.", "Bekijk hooks, topics, formats, ideale lengte, visual style en evidence rows na data-import of scenarioselectie."],
            ["Calendar", "Plant gereviewde concepten in een publicatieschema.", "Controleer status, platform, data en titels. Items uit Draft Studio zijn bedoeld voor review, niet voor auto-publish."],
            ["Content Library", "Bewaart concepten, outlines en herbruikbare assets.", "Bekijk status, doelgroep, contenttype en volgende actie. Opgeslagen drafts uit Research verschijnen hier."],
            ["Reports", "Vat prestatiesignalen samen voor presentatie en submission.", "Vergelijk platformen, median engagement, sample size, best signal, coverage en exporteer het evidence pack."],
            ["Saved Searches", "Bewaart herbruikbare onderzoeksscenario's.", "Gebruik combinaties van ICP, platform, objective en tone om een discovery flow te herhalen."],
            ["Brand Voice", "Bevat schrijfrichtlijnen voor Pro Actief en Amsterdam Game Lab.", "Controleer geloofwaardigheid, tone-grenzen, bewijsvereisten en contentprincipes voor publicatie."],
            ["Team", "Toont ownership en reviewverantwoordelijkheden.", "Gebruik dit om te laten zien wie onderzoekt, schrijft, reviewt en goedkeurt."],
            ["Settings", "Beheert taal, opslag, publishing safety en backup.", "Wissel taal, zet verificatie-safeguards aan, exporteer backup en wis lokale workspace-data indien nodig."],
            ["Update Profile", "Bewerkt het workspace-profiel in navbar en profielmenu.", "Wijzig naam, rol, e-mail en profielfoto voor demo of presentatie."],
            ["Data Sources", "Verklaart en beheert geimporteerd bewijs.", "Importeer CSV, bekijk geaccepteerde bronnen, reset demo scenario en exporteer backupdata."],
            ["Shortcuts", "Geeft snelle routes naar veelgebruikte workflows.", "Spring naar Research, Content Library, Reports of herhaal de demo flow."],
        ],
        "buttons": [
            ["Login", "Start de app en opent Overview."],
            ["Back to landing", "Keert terug van het loginformulier naar de landing page."],
            ["Import CSV", "Voegt onderzoeksrijen toe uit platformexports of handmatig onderzoek."],
            ["Filters", "Opent de drawer met actieve scenariofilters."],
            ["Save search", "Bewaart de huidige filtercombinatie."],
            ["Generate draft", "Maakt een concept op basis van actieve patronen, geselecteerde rij, ICP, platform, objective en tone."],
            ["Save draft", "Bewaart het huidige concept in Content Library."],
            ["Add to Calendar", "Maakt een planningitem voor review."],
            ["Export analysis / Export report / Export backup", "Downloadt JSON evidence voor reproduceerbaarheid en hackathon submission."],
            ["Reset demo scenario", "Herstelt de ingebouwde demodata voor presentatiegebruik."],
            ["Clear workspace", "Verwijdert lokale browserdata voor een schone demo."],
            ["Logout", "Keert terug naar de landing page."],
        ],
        "workflow": [
            "Contentonderzoek: kies een scenario, importeer bewijs, rank rijen en selecteer een sterke referentie.",
            "Patroonextractie: AGLI vat hook-, topic-, format-, lengte- en visual-style patronen samen uit de actieve evidence set.",
            "Conceptcreatie: Draft Studio combineert het gekozen patroon met Pro Actief guardrails en maakt copy, hooks, video outlines of hashtags.",
            "Menselijke review: het team controleert feitelijke juistheid, brand fit, juridisch risico en bewijskwaliteit.",
            "Planning en rapportage: save draft, add calendar item, export reports en documenteer wat is geverifieerd.",
        ],
        "csv": [
            ["title", "Contenttitel of hook."],
            ["platform", "LinkedIn, YouTube, Instagram of TikTok."],
            ["account", "Naam van het bronaccount."],
            ["type", "Text post, carousel, video, article of vergelijkbaar format."],
            ["views", "Geverifieerde of geimporteerde view count."],
            ["engagements", "Likes, comments, shares of platform engagement count."],
            ["date", "Publicatiedatum."],
            ["hookType / topic / visualStyle", "Optionele patroonvelden die Draft Studio verbeteren."],
            ["verificationNote", "Optionele menselijke notitie over bronkwaliteit."],
        ],
        "demo_script": [
            "Start op de landing page en leg het probleem uit: serious games zijn category-creation, dus content moet aandacht verdienen.",
            "Login en toon Overview als command center.",
            "Open Research, wijzig scenario-dropdowns en selecteer een goed presterende rij.",
            "Toon Pattern Insights en leg uit hoe de app hooks, topics, formats, length en visual style vindt.",
            "Genereer Draft Studio output, sla deze op en voeg toe aan Calendar.",
            "Open Reports en exporteer het evidence pack.",
            "Sluit af met guardrails: geen auto-publish, geen verzonnen claims, en elk concept vereist menselijke review.",
        ],
        "troubleshooting": [
            ["Pagina verandert na refresh", "Refresh houdt nu de huidige pagina vast. Alleen login opent bewust Overview."],
            ["Dropdown blijft open", "Klik buiten de dropdown om deze te sluiten."],
            ["Drawer blijft open", "Klik buiten de rechter drawer of gebruik het X-icoon."],
            ["Concept past niet", "Wijzig ICP, platform, objective, tone, geselecteerde rij of content pillar en genereer opnieuw."],
            ["Metrics zijn niet verified", "Importeer officiele platformexports of voeg verification notes toe voordat je extern gebruikt."],
        ],
    },
}


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |")
    return "\n".join(lines)


def write_markdown():
    parts = [
        "# AGLI Manual Book",
        "",
        "Amsterdam Game Lab Intelligence",
        "",
        "This manual is generated for hackathon submission and presentation support.",
        "",
    ]
    for language, data in manual.items():
        parts.extend([f"# {language}", "", f"## {data['subtitle']}", ""])
        parts.extend(data["intro"])
        parts.append("")
        parts.append("## Core Principles")
        parts.extend([f"- {item}" for item in data["principles"]])
        parts.append("")
        parts.append("## Quick Start")
        parts.extend([f"{index}. {item}" for index, item in enumerate(data["quick_start"], start=1)])
        parts.append("")
        parts.append("## Global Controls")
        parts.append(markdown_table(["Control", "Function"], data["global"]))
        parts.append("")
        parts.append("## Page Guide")
        parts.append(markdown_table(["Page", "Purpose", "Main actions"], data["pages"]))
        parts.append("")
        parts.append("## Button and Action Reference")
        parts.append(markdown_table(["Button / Action", "What it does"], data["buttons"]))
        parts.append("")
        parts.append("## Recommended Workflow")
        parts.extend([f"- {item}" for item in data["workflow"]])
        parts.append("")
        parts.append("## CSV Data Guide")
        parts.append(markdown_table(["Column", "Meaning"], data["csv"]))
        parts.append("")
        parts.append("## Presentation Script")
        parts.extend([f"{index}. {item}" for index, item in enumerate(data["demo_script"], start=1)])
        parts.append("")
        parts.append("## Troubleshooting")
        parts.append(markdown_table(["Issue", "Fix"], data["troubleshooting"]))
        parts.append("")
    OUT_MD.write_text("\n".join(parts), encoding="utf-8")


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 117, 103)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    document.add_paragraph("")


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def write_docx():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(3, 44, 54)

    document.add_heading("AGLI Manual Book", level=0)
    document.add_paragraph("Amsterdam Game Lab Intelligence")
    document.add_paragraph("Generated for hackathon submission, tutorial support, and presentation preparation.")

    meta = [
        ["Project", "AGLI - Amsterdam Game Lab Intelligence"],
        ["Audience", "HR, L&D, team leads, managers, content strategists, and hackathon reviewers"],
        ["Scope", "Landing, login, app pages, controls, workflows, data handling, and presentation script"],
        ["Safety", "No auto-publishing. No fabricated metrics, claims, quotes, or case studies."],
    ]
    add_table(document, ["Field", "Detail"], meta)

    for language, data in manual.items():
        document.add_page_break()
        document.add_heading(language, level=1)
        document.add_heading(data["subtitle"], level=2)
        for paragraph in data["intro"]:
            document.add_paragraph(paragraph)

        document.add_heading("Core Principles", level=2)
        add_bullets(document, data["principles"])

        document.add_heading("Quick Start", level=2)
        add_numbered(document, data["quick_start"])

        document.add_heading("Global Controls", level=2)
        add_table(document, ["Control", "Function"], data["global"])

        document.add_heading("Page Guide", level=2)
        add_table(document, ["Page", "Purpose", "Main actions"], data["pages"])

        document.add_heading("Button and Action Reference", level=2)
        add_table(document, ["Button / Action", "What it does"], data["buttons"])

        document.add_heading("Recommended Workflow", level=2)
        add_bullets(document, data["workflow"])

        document.add_heading("CSV Data Guide", level=2)
        add_table(document, ["Column", "Meaning"], data["csv"])

        document.add_heading("Presentation Script", level=2)
        add_numbered(document, data["demo_script"])

        document.add_heading("Troubleshooting", level=2)
        add_table(document, ["Issue", "Fix"], data["troubleshooting"])

    document.save(OUT_DOCX)


def main():
    DOCS.mkdir(exist_ok=True)
    write_markdown()
    write_docx()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
